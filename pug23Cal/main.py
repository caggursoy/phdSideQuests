# -*- coding: utf-8 -*-
"""
Created on Tue Jun  6 11:32:05 2023

@author: Cagatay.Guersoy
"""
import urllib.request, vobject, os, json, sys, PyPDF2
from gcsa.google_calendar import GoogleCalendar
from gcsa.event import Event
from datetime import datetime, timezone, timedelta
import pytz
from tzlocal import get_localzone # $ pip install tzlocal
import ssl
from tabula import read_pdf
from pathlib import Path
from docx import Document

#from O365 import Account, MSGraphProtocol

# define main
def main():
    del_flag = True
    print('Running program')
    gcal = conf_gcal()
    print('Connected to the Google Calendar successfully...')
    if del_flag:
        flush_cal(gcal)
        print('Flushed the calendar!!!')
    read_and_create(gcal)
    print('Created the calendar (if new events existed)')
    
    
def conf_gcal():
    with open('./credentials.json', 'r') as f:
        credentials_data = json.load(f)
    calendar_id = credentials_data['installed']['calendar_id']
    print(calendar_id)
    try:
        gcal = GoogleCalendar(calendar_id, credentials_path = './credentials.json') # set the gcal
    except Exception as e:
        if os.path.isfile('token.pickle'):
            os.remove('token.pickle') # remove the pickle file
            gcal = GoogleCalendar(calendar_id, credentials_path = './credentials.json') # try to set the gcal again
        else:
            print(e)
            sys.exit()
    # print(gcal)
    # gcal_events = [] # allocate empty list for gcal events
    # for g_ev in gcal: # in a loop save every event to a list
    #     gcal_events.append(g_ev)
    # print(gcal_events)
    return(gcal)
    

def read_and_create(calendar):
    # First download the pdf files
    main_filename = ['https://pug2023.de/wp-content/uploads/2023/06/PuG23_Program_Booklet.pdf']
    pdf_filenames = ['https://pug2023.de/wp-content/uploads/2023/05/'+day for day in ['Thursday_Program.pdf', 'Friday_Program-1.pdf', 'Saturday_Program.pdf']]
    saved_filenames = []
    days = ['08.06.2023', '09.06.2023', '10.06.2023']
    for file in main_filename + pdf_filenames:
        #urllib.request.urlretrieve(file, str(Path.cwd() / file[file.rfind('/')+1:]))
        saved_filenames.append(str(Path.cwd() / (file[file.rfind('/')+1:file.rfind('.pdf')]+'.docx')))
    # print(saved_filenames)
    event_list = []
    for j in range(1, len(saved_filenames)):
        if j != 0:
            filename = saved_filenames[j] # to be changed later to j
            # print('File:', filename)
            # get doc
            document = Document(filename)
            # create pages list
            line_list = []
            for table in document.tables:
            # table = document.tables[0]
                for i, row in enumerate(table.rows):
                    text = (cell.text for cell in row.cells)
                    aux_list = list(text)
                    # print(aux_list)
                    line_list.append(aux_list)
                    # print(line_list)
                    for q in range(len(line_list)):
                        if 'Symposia Session' not in '\t'.join(line_list[q]) and 'Keynote' not in '\t'.join(line_list[q]) and 'Social evening' not in '\t'.join(line_list[q]):
                            curr_line = line_list[q]
                            curr_line = [el.replace('\n','') for el in curr_line]
                            # print(curr_line)
                            times = [curr_line[0][:curr_line[0].find('-')].replace(' ',''), curr_line[0][curr_line[0].find('-')+1:]]
                            # print('Times:',times)
                            place = curr_line[1]
                            # print('Place:',place)
                            head = curr_line[2].split('Chairs: ')[0]
                            # print('Header:', head)
                            chairs = curr_line[2].split('Chairs: ')[1].replace('Chairs: ','').split(',')
                            # print('Chairs:', chairs)
                            start_time = datetime(int(days[j-1].split('.')[2]),  # days variable to be changed later to [q]
                                                   int(days[j-1].split('.')[1]),
                                                   int(days[j-1].split('.')[0]),
                                                   int(times[0].split(':')[0]),
                                                   int(times[0].split(':')[1])
                                                   )
                            end_time = datetime(int(days[j-1].split('.')[2]),  # days variable to be changed later to [q]
                                                         int(days[j-1].split('.')[1]),
                                                         int(days[j-1].split('.')[0]),
                                                         int(times[1].split(':')[0]),
                                                         int(times[1].split(':')[1])
                                                         )
                            event = Event(
                                summary = head,
                                description = "".join(str(x) for x in ['Chairs: '] + [chair +',' for chair in chairs]),
                                start=start_time,
                                end=end_time,
                                location=place,
                                minutes_before_popup_reminder = 10,
                            )
                            # print(event)
                            if event not in event_list:
                                event_list.append(event)
                                
                        elif 'Keynote' in '\t'.join(line_list[q]):
                            curr_line = line_list[q]
                            # print(curr_line)
                            times = curr_line[0].split(' -\n')
                            # print('Times:',times)
                            place = curr_line[1]
                            # print('Place:',place)
                            head = curr_line[2].split('\n')[0]
                            # print('Header:', head)
                            chairs = curr_line[2].split('\n')[1]
                            start_time = datetime(int(days[j-1].split('.')[2]),  # days variable to be changed later to [q]
                                                   int(days[j-1].split('.')[1]),
                                                   int(days[j-1].split('.')[0]),
                                                   int(times[0].split(':')[0]),
                                                   int(times[0].split(':')[1])
                                                   )
                            end_time = datetime(int(days[j-1].split('.')[2]),  # days variable to be changed later to [q]
                                                         int(days[j-1].split('.')[1]),
                                                         int(days[j-1].split('.')[0]),
                                                         int(times[1].split(':')[0]),
                                                         int(times[1].split(':')[1])
                                                         )
                            event = Event(
                                summary = head,
                                description = 'From: ' + chairs[0],
                                start=start_time,
                                end=end_time,
                                location=place,
                                minutes_before_popup_reminder = 10,
                            )
                            # print(event)
                            calendar.add_event(event)
                            event_list.append(event)
    
    # Now compare and add events
    gcal_events = []
    for g_ev in calendar:
        gcal_events.append(g_ev)
    # print('Existing:',gcal_events)
    # print('New:',event_list)
    # print(set([ev_list.summary for ev_list in event_list]).difference(set([gcal_ev.summary for gcal_ev in gcal_events])))
    for ev in set([ev_list.summary for ev_list in event_list]).difference(set([gcal_ev.summary for gcal_ev in gcal_events])):
        # calendar.add_event(ev_list[ev_list.index(ev)])
        calendar.add_event(event_list[[ev_list.summary for ev_list in event_list].index(ev)])
                            
                            
    # print(event_list, len(event_list))
    
def flush_cal(gcal):
    for g_ev in gcal: # in a loop save every event to a list
        gcal.delete_event(g_ev)
                            

# run main
if __name__ == "__main__":
    main()
    